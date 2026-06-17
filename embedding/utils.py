import csv
import os
import pdfplumber
import nltk
from html.parser import HTMLParser
from typing import Optional


nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)


def get_file_paths(root_dir: str, file_extensions: list[str]) -> list[str]:
    """
    Retrieves a list of paths to all files with specified extensions in the given root directory and its subdirectories.

    Args:
        root_dir (str): The root directory to search for files.
        file_extensions (list[str]): A list of file extensions to retrieve. For example, ["txt", "pdf"]

    Returns:
        List[str]: A list of file paths to all matching files found within the root directory and its subdirectories.
    """
    file_paths = []
    
    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            if any(filename.endswith(f".{ext}") for ext in file_extensions):
                file_paths.append(os.path.join(dirpath, filename))
    
    return file_paths



def read_text_file(file_path: str) -> str:
    """
    Reads the content of a text file and returns it as a single string.

    Args:
        file_path (str): The path to the .txt file to read.

    Returns:
        str: The content of the file as a single string.
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    return content


def read_pdf_file(file_path: str) -> str:
    """Read the content of a PDF file and return it as a single string.

    Both prose text and tables are extracted.  Tables are formatted as
    pipe-delimited rows so their structure is preserved in the output
    rather than being collapsed into a single garbled line.

    Args:
        file_path (str): The path to the PDF file to read.

    Returns:
        str: The full content of the PDF as a single string.
    """
    parts: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = []
            try:
                tables = page.find_tables()
            except Exception:
                pass

            if not tables:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(page_text)
                continue

            # Extract prose text outside detected table regions.
            text_page = page
            for table in tables:
                try:
                    text_page = text_page.outside_bbox(table.bbox)
                except Exception:
                    pass
            prose = text_page.extract_text() or ""
            if prose.strip():
                parts.append(prose)

            # Append each table as pipe-delimited structured text.
            for table in tables:
                try:
                    rows = table.extract()
                    table_text = _format_table_as_text(rows)
                    if table_text:
                        parts.append(table_text)
                except Exception:
                    pass

    return "\n".join(parts)


def split_text_into_sentences(text: str, language: str) -> list[str]:
    """
    Splits the given text into a list of sentences using NLTK's sentence tokenizer.

    Args:
        text (str): The input text to split into sentences.
        language (str): The language of the text for the sentence tokenizer

    Returns:
        list[str]: A list of sentences.
    """
    sentences = nltk.sent_tokenize(text, language=language)
    return sentences


def chunk_sentences(sentences: list[str], chunk_size: int, overlap_size: int) -> list[str]:
    """
    Groups a list of sentences into overlapping chunks.

    Args:
        sentences (list[str]): A list of sentences to be chunked.
        chunk_size (int): The number of sentences in each chunk.
        overlap_size (int): The number of sentences to overlap between consecutive chunks.

    Returns:
        list[str]: A list of text chunks with the specified overlap.
        
    Raises:
        ValueError: If overlap_size is greater than or equal to chunk_size.
    """
    if overlap_size >= chunk_size:
        raise ValueError("overlap_size must be smaller than chunk_size.")

    chunks = []
    # The step of the range is the distance between the start of consecutive chunks.
    step = chunk_size - overlap_size
    for i in range(0, len(sentences), step):
        chunk = " ".join(sentences[i:i + chunk_size])
        chunks.append(chunk)
    return chunks


# ---------------------------------------------------------------------------
# HTML helper (used by both read_html_file and read_markdown_file)
# ---------------------------------------------------------------------------

class _TextExtractor(HTMLParser):
    """Minimal HTMLParser subclass that collects visible text nodes."""

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_tags = {"script", "style", "head"}
        self._skip = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._skip_tags:
            self._skip = True

    def handle_endtag(self, tag: str) -> None:
        if tag in self._skip_tags:
            self._skip = False
        # Insert a space after block-level tags to prevent words merging.
        if tag in {"p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6",
                   "tr", "br"}:
            self._parts.append(" ")

    def handle_data(self, data: str) -> None:
        if not self._skip:
            self._parts.append(data)

    def get_text(self) -> str:
        return " ".join("".join(self._parts).split())


def _html_to_text(html: str) -> str:
    """Strip HTML tags and return clean visible text.

    Args:
        html: Raw HTML string.

    Returns:
        Plain text with whitespace normalised.
    """
    extractor = _TextExtractor()
    extractor.feed(html)
    return extractor.get_text()


# ---------------------------------------------------------------------------
# New file-type readers
# ---------------------------------------------------------------------------

def read_docx_file(file_path: str) -> str:
    """Read a Word (.docx) document and return its text as a single string.

    Paragraphs are joined with newlines.  Empty paragraphs (used in Word
    for visual spacing) are omitted.

    Args:
        file_path: Path to the .docx file.

    Returns:
        The document's text content.
    """
    import docx  # python-docx; imported here to keep the module importable
                  # even when the package is not installed.

    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def read_html_file(file_path: str, encoding: str = "utf-8") -> str:
    """Read an HTML file and return its visible text content.

    Tags, scripts, and style blocks are stripped.  Whitespace is normalised
    to single spaces.

    Args:
        file_path: Path to the .html (or .htm) file.
        encoding: File encoding (default ``utf-8``).

    Returns:
        Plain text extracted from the HTML.
    """
    with open(file_path, "r", encoding=encoding, errors="replace") as fh:
        html = fh.read()
    return _html_to_text(html)


def read_markdown_file(file_path: str, encoding: str = "utf-8") -> str:
    """Read a Markdown file and return its plain-text content.

    The file is rendered to HTML via ``markdown-it-py`` (already a project
    dependency) and then stripped of tags so the text passed to the
    sentence tokeniser is clean prose rather than raw Markdown syntax.

    Args:
        file_path: Path to the .md file.
        encoding: File encoding (default ``utf-8``).

    Returns:
        Plain text extracted from the Markdown source.
    """
    from markdown_it import MarkdownIt  # markdown-it-py is already installed.

    with open(file_path, "r", encoding=encoding, errors="replace") as fh:
        source = fh.read()

    md = MarkdownIt()
    html = md.render(source)
    return _html_to_text(html)


def read_csv_file(
    file_path: str,
    encoding: str = "utf-8",
    delimiter: str = ",",
    has_header: bool = True,
    max_rows: Optional[int] = None,
) -> str:
    """Read a CSV file and return its contents as human-readable prose.

    Each data row is serialised as ``"column: value, column: value."`` so
    that the sentence tokeniser can treat individual rows as discrete units.

    Args:
        file_path: Path to the .csv file.
        encoding: File encoding (default ``utf-8``).
        delimiter: Column delimiter (default ``,``).
        has_header: When ``True`` (default) the first row is treated as
            column names.  When ``False`` columns are labelled ``col_0``,
            ``col_1``, etc.
        max_rows: Optional cap on the number of data rows read.  ``None``
            reads all rows.

    Returns:
        A single string with one sentence per row, ready for sentence
        tokenisation.
    """
    rows: list[str] = []

    with open(file_path, "r", encoding=encoding, errors="replace",
              newline="") as fh:
        if has_header:
            reader = csv.DictReader(fh, delimiter=delimiter)
            for idx, row in enumerate(reader):
                if max_rows is not None and idx >= max_rows:
                    break
                pairs = ", ".join(
                    f"{key}: {val.strip()}"
                    for key, val in row.items()
                    if val and val.strip()
                )
                rows.append(pairs + ".")
        else:
            plain = csv.reader(fh, delimiter=delimiter)
            for idx, row in enumerate(plain):
                if max_rows is not None and idx >= max_rows:
                    break
                pairs = ", ".join(
                    f"col_{i}: {val.strip()}" for i, val in enumerate(row)
                )
                rows.append(pairs + ".")

    return "\n".join(rows)


# ---------------------------------------------------------------------------
# Section-aware HTML parser (used by read_document_sentences)
# ---------------------------------------------------------------------------

class _SectionHTMLParser(HTMLParser):
    """HTML parser that tracks h1-h6 headings and collects text blocks.

    Yields ``(section_title, block_text)`` pairs where ``section_title`` is
    the text of the most recent heading encountered, or ``None`` if no
    heading has been seen yet.
    """

    _HEADING_TAGS = frozenset({"h1", "h2", "h3", "h4", "h5", "h6"})
    _SKIP_TAGS = frozenset({"script", "style", "head"})
    _BLOCK_TAGS = frozenset({
        "p", "div", "li", "td", "th", "article", "section",
        "blockquote", "figcaption", "br",
    })

    def __init__(self) -> None:
        super().__init__()
        self._current_section: Optional[str] = None
        self._in_heading: bool = False
        self._heading_buf: list[str] = []
        self._block_buf: list[str] = []
        self._skip: bool = False
        self.blocks: list[tuple[Optional[str], str]] = []

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._SKIP_TAGS:
            self._skip = True
        elif tag in self._HEADING_TAGS:
            self._flush_block()
            self._in_heading = True
            self._heading_buf = []
        elif tag in self._BLOCK_TAGS:
            self._flush_block()

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS:
            self._skip = False
        elif tag in self._HEADING_TAGS:
            heading_text = " ".join(self._heading_buf).strip()
            if heading_text:
                self._current_section = heading_text
            self._in_heading = False
        elif tag in self._BLOCK_TAGS:
            self._flush_block()

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        if self._in_heading:
            self._heading_buf.append(data)
        else:
            self._block_buf.append(data)

    def _flush_block(self) -> None:
        text = " ".join(self._block_buf).strip()
        if text:
            self.blocks.append((self._current_section, text))
        self._block_buf = []

    def finalize(self) -> list[tuple[Optional[str], str]]:
        """Flush any remaining buffered text and return all blocks."""
        self._flush_block()
        return self.blocks


# ---------------------------------------------------------------------------
# Private per-format sentence extractors with metadata
# ---------------------------------------------------------------------------

def _make_sent(
    text: str,
    page_number: Optional[int] = None,
    section_title: Optional[str] = None,
) -> dict:
    """Build a sentence metadata dict.

    Only non-None values are included so that ChromaDB, which requires
    string/int/float/bool metadata values, never sees ``None``.
    """
    entry: dict = {"text": text}
    if page_number is not None:
        entry["page_number"] = page_number
    if section_title is not None:
        entry["section_title"] = section_title
    return entry


def _sentences_txt(file_path: str, language: str) -> list[dict]:
    """Extract sentences from a plain-text file."""
    text = read_text_file(file_path)
    return [
        _make_sent(s)
        for s in nltk.sent_tokenize(text, language=language)
        if s.strip()
    ]


def _format_table_as_text(rows: list[list]) -> str:
    """Format a pdfplumber table as a pipe-delimited markdown-style table.

    The first row is treated as the header.  Entirely empty rows are
    dropped.  A separator line is inserted below the header when more
    than one row is present so the output resembles a Markdown table,
    making it readable in both prose and structured contexts.

    Args:
        rows: A list of rows; each row is a list of cell values (str or None).

    Returns:
        A formatted string, or an empty string if the table has no usable
        content.
    """
    if not rows:
        return ""

    normalised = [
        [str(cell).strip() if cell is not None else "" for cell in row]
        for row in rows
        if any(cell is not None and str(cell).strip() for cell in row)
    ]
    if not normalised:
        return ""

    formatted_rows = ["| " + " | ".join(row) + " |" for row in normalised]

    if len(formatted_rows) > 1:
        n_cols = len(normalised[0])
        separator = "| " + " | ".join(["---"] * n_cols) + " |"
        formatted_rows.insert(1, separator)

    return "\n".join(formatted_rows)


def _extract_pdf_page_sentences(
    page, page_num: int, language: str
) -> list[dict]:
    """Extract sentences and tables from a single PDF page.

    Tables are detected with pdfplumber layout analysis and formatted as
    pipe-delimited text to preserve row/column structure.  Non-table prose
    is sentence-tokenised normally.  Tables always follow prose in the
    returned list, which is sufficient for RAG retrieval since each chunk
    is retrieved independently.

    Args:
        page: A pdfplumber.Page object.
        page_num: 1-based page index stored as page_number metadata.
        language: NLTK sentence-tokeniser language.

    Returns:
        A list of sentence/table dicts with page_number set on every entry.
    """
    result: list[dict] = []

    try:
        tables = page.find_tables()
    except Exception:
        tables = []

    if not tables:
        page_text = page.extract_text() or ""
        for sent in nltk.sent_tokenize(page_text, language=language):
            if sent.strip():
                result.append(_make_sent(sent, page_number=page_num))
        return result

    # Exclude table bounding boxes so cell text is not duplicated in prose.
    text_page = page
    for table in tables:
        try:
            text_page = text_page.outside_bbox(table.bbox)
        except Exception:
            pass

    prose_text = text_page.extract_text() or ""
    for sent in nltk.sent_tokenize(prose_text, language=language):
        if sent.strip():
            result.append(_make_sent(sent, page_number=page_num))

    # Append each table as a single structured block.
    for table in tables:
        try:
            rows = table.extract()
            table_text = _format_table_as_text(rows)
            if table_text:
                result.append(_make_sent(table_text, page_number=page_num))
        except Exception:
            pass

    return result


def _sentences_pdf(file_path: str, language: str) -> list[dict]:
    """Extract sentences from a PDF, tagging each with its page number.

    Tables on each page are detected and formatted as pipe-delimited rows
    to preserve their structure for retrieval.  Falls back to plain text
    extraction per page on any per-page error.

    Args:
        file_path: Path to the PDF file.
        language: NLTK sentence-tokeniser language.

    Returns:
        List of sentence/table dicts with page_number metadata.
    """
    result: list[dict] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    result.extend(
                        _extract_pdf_page_sentences(page, page_num, language)
                    )
                except Exception:
                    try:
                        page_text = page.extract_text() or ""
                        for sent in nltk.sent_tokenize(page_text, language=language):
                            if sent.strip():
                                result.append(
                                    _make_sent(sent, page_number=page_num)
                                )
                    except Exception:
                        pass
    except Exception as exc:
        print(f"  Warning: could not fully read PDF '{file_path}': {exc}")
    return result


def _sentences_docx(file_path: str, language: str) -> list[dict]:
    """Extract sentences from a Word document, tagging with heading sections."""
    import docx  # python-docx

    result: list[dict] = []
    current_section: Optional[str] = None
    try:
        document = docx.Document(file_path)
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                current_section = text
            else:
                for sent in nltk.sent_tokenize(text, language=language):
                    if sent.strip():
                        result.append(
                            _make_sent(sent, section_title=current_section)
                        )
    except Exception as exc:
        print(f"  Warning: could not fully read DOCX '{file_path}': {exc}")
    return result


def _sentences_html(file_path: str, language: str) -> list[dict]:
    """Extract sentences from an HTML file, tagging with section headings."""
    result: list[dict] = []
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            html = fh.read()
        parser = _SectionHTMLParser()
        parser.feed(html)
        blocks = parser.finalize()
        for section_title, block_text in blocks:
            for sent in nltk.sent_tokenize(block_text, language=language):
                if sent.strip():
                    result.append(_make_sent(sent, section_title=section_title))
    except Exception as exc:
        print(f"  Warning: could not fully read HTML '{file_path}': {exc}")
    return result


def _sentences_md(file_path: str, language: str) -> list[dict]:
    """Extract sentences from a Markdown file, tagging with section headings."""
    from markdown_it import MarkdownIt

    result: list[dict] = []
    current_section: Optional[str] = None
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            source = fh.read()
        tokens = MarkdownIt().parse(source)
        in_heading = False
        for tok in tokens:
            if tok.type == "heading_open":
                in_heading = True
            elif tok.type == "heading_close":
                in_heading = False
            elif tok.type == "inline":
                content = tok.content.strip()
                if not content:
                    continue
                if in_heading:
                    # Strip any remaining inline markdown syntax (e.g. **bold**)
                    current_section = _html_to_text(
                        MarkdownIt().render(content)
                    ).strip() or content
                else:
                    for sent in nltk.sent_tokenize(content, language=language):
                        if sent.strip():
                            result.append(
                                _make_sent(sent, section_title=current_section)
                            )
    except Exception as exc:
        print(f"  Warning: could not fully read Markdown '{file_path}': {exc}")
    return result


def _sentences_csv(file_path: str) -> list[dict]:
    """Return each non-empty CSV row as a sentence dict (no page/section)."""
    result: list[dict] = []
    try:
        text = read_csv_file(file_path)
        for line in text.splitlines():
            if line.strip():
                result.append(_make_sent(line.strip()))
    except Exception as exc:
        print(f"  Warning: could not fully read CSV '{file_path}': {exc}")
    return result


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

_EXTENSION_MAP = {
    ".txt": _sentences_txt,
    ".pdf": _sentences_pdf,
    ".docx": _sentences_docx,
    ".html": _sentences_html,
    ".htm": _sentences_html,
    ".md": _sentences_md,
    ".csv": _sentences_csv,
}


def read_document_sentences(
    file_path: str,
    language: str = "english",
) -> list[dict]:
    """Read a document and return its sentences as metadata-enriched dicts.

    Each dict always contains ``"text"`` and optionally ``"page_number"``
    (int, PDF only) and ``"section_title"`` (str, for structured formats).
    Keys are omitted entirely when the value is not available, so callers
    and ChromaDB never receive ``None`` values.

    Supported extensions: ``.txt``, ``.pdf``, ``.docx``, ``.html``,
    ``.htm``, ``.md``, ``.csv``.

    Args:
        file_path: Absolute path to the source document.
        language: NLTK sentence-tokeniser language (default ``"english"``).

    Returns:
        A list of sentence dicts.  Returns an empty list if the file cannot
        be read or produces no extractable text.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = os.path.splitext(file_path)[1].lower()
    extractor = _EXTENSION_MAP.get(ext)
    if extractor is None:
        raise ValueError(
            f"Unsupported file extension '{ext}'. "
            f"Supported: {sorted(_EXTENSION_MAP)}."
        )
    if ext in (".txt", ".pdf", ".docx", ".html", ".htm", ".md"):
        return extractor(file_path, language)
    # CSV has no language parameter
    return extractor(file_path)

